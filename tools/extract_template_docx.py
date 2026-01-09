import json
import re
from pathlib import Path

from docx import Document

BLOCK_RE = re.compile(r"^Bloque\s+([A-E])\s+\u2013\s+(.*?)\s+\((\d+)%\)\s*$", re.IGNORECASE)


def clean(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def extract(docx_path: str, base_code: str):
    doc = Document(docx_path)

    # 1) Capture paragraphs to detect block headers.
    paragraphs = [clean(p.text) for p in doc.paragraphs if clean(p.text)]
    block_headers = []
    for idx, t in enumerate(paragraphs):
        m = BLOCK_RE.match(t)
        if m:
            block_headers.append((m.group(1).upper(), m.group(2), int(m.group(3)), t))

    # 2) Assign tables to blocks by order.
    blocks = []
    tables = doc.tables

    for i, (letter, title, pct, raw) in enumerate(block_headers):
        if i >= len(tables):
            break
        tbl = tables[i]
        rows = []
        for r, row in enumerate(tbl.rows):
            cells = [clean(c.text) for c in row.cells]
            # Skip typical header row.
            if r == 0 and any("Subcriterio" in c for c in cells):
                continue
            sub = cells[0] if len(cells) > 0 else ""
            desc = cells[1] if len(cells) > 1 else ""
            if not sub and not desc:
                continue
            rows.append(
                {
                    "subcriterion": sub,
                    "description": desc,
                    "scale": "1-5",
                }
            )

        blocks.append(
            {
                "code": letter,
                "title": title,
                "weight_percent": pct,
                "items": rows,
            }
        )

    data = {
        "base_code": base_code,
        "source_file": Path(docx_path).name,
        "blocks": blocks,
        "notes": {
            "expected_blocks": ["A", "B", "C", "D", "E"],
            "observed_blocks": [b["code"] for b in blocks],
        },
    }
    return data


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 4:
        raise SystemExit(
            "Usage: python tools/extract_template_docx.py <docx_path> <base_code> <json_output_path>"
        )

    docx_path = sys.argv[1]
    base_code = sys.argv[2]
    out_path = sys.argv[3]

    payload = extract(docx_path, base_code=base_code)
    Path(out_path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print("OK ->", out_path)
