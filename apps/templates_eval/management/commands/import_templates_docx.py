import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Tuple

from django.core.management.base import BaseCommand, CommandError
from django.db import transaction

from apps.templates_eval.models import EvaluationTemplate, TemplateSection, TemplateQuestion

try:
    from docx import Document  # python-docx
except Exception as e:  # pragma: no cover
    Document = None


# Ajusta estos literales a los que uses en tu modelo si difieren
TYPE_SCALE = "SCALE_1_5"
TYPE_YESNO = "YES_NO"
TYPE_TEXT = "TEXT"

REQ_PATTERNS = [
    re.compile(r"\[REQ\]", re.IGNORECASE),
    re.compile(r"\bREQ\b", re.IGNORECASE),
    re.compile(r"\*$"),
]

TYPE_RULES: List[Tuple[re.Pattern, str]] = [
    (re.compile(r"\(1\s*-\s*5\)", re.IGNORECASE), TYPE_SCALE),
    (re.compile(r"\[SCALE\]", re.IGNORECASE), TYPE_SCALE),
    (re.compile(r"\(Y\s*/\s*N\)", re.IGNORECASE), TYPE_YESNO),
    (re.compile(r"\[YESNO\]", re.IGNORECASE), TYPE_YESNO),
    (re.compile(r"\(TEXT\)", re.IGNORECASE), TYPE_TEXT),
    (re.compile(r"\[TEXT\]", re.IGNORECASE), TYPE_TEXT),
]

ADMIN_FIELD_PREFIXES = (
    "nombre", "apellidos", "departamento", "puesto", "grupo profesional",
    "fecha", "periodo", "período", "firma", "conclusión", "conclusion",
    "suma", "nivel global"
)

HEADER_KEYWORDS = (
    "criterio", "factor", "peso", "puntuación", "puntuacion", "bloque",
    "evidencias", "objetivos", "validación", "validacion"
)


def clean_cell_text(s: str) -> str:
    return " ".join((s or "").replace("\n", " ").split()).strip()


def looks_like_admin_field(s: str) -> bool:
    t = (s or "").strip().lower()
    return any(t.startswith(p) for p in ADMIN_FIELD_PREFIXES)


def looks_like_header_row(cells: List[str]) -> bool:
    joined = " ".join(cells).lower()
    return any(k in joined for k in HEADER_KEYWORDS)


def extract_weight(cells: List[str]) -> Optional[str]:
    # Devuelve el primer "NN%" que encuentre
    for c in cells:
        m = re.search(r"(\d{1,3})\s*%", c)
        if m:
            return f"{m.group(1)}%"
    return None


def derive_base_code_from_filename(path: Path) -> str:
    m = re.match(r"^(P\d{2})_", path.stem, flags=re.IGNORECASE)
    if m:
        return m.group(1).upper()
    return path.stem[:20].upper()


def template_fingerprint(parsed: "ParsedTemplate") -> str:
    payload = [parsed.name]
    for sec in parsed.sections:
        payload.append(f"[S]{sec.name}")
        for q in sec.questions:
            payload.append(f"[Q]{q.question_type}|{int(q.is_required)}|{q.text}")
    raw = "\n".join(payload).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def normalize_question_text(raw: str) -> str:
    t = raw.strip()
    # elimina marcadores conocidos
    t = re.sub(r"\[REQ\]", "", t, flags=re.IGNORECASE).strip()
    t = re.sub(r"\[SCALE\]|\[YESNO\]|\[TEXT\]", "", t, flags=re.IGNORECASE).strip()
    t = re.sub(r"\(1\s*-\s*5\)|\(Y\s*/\s*N\)|\(TEXT\)", "", t, flags=re.IGNORECASE).strip()
    # quita asterisco final
    t = re.sub(r"\*$", "", t).strip()
    return t


def detect_is_required(raw: str) -> bool:
    s = raw.strip()
    return any(p.search(s) for p in REQ_PATTERNS)


def detect_question_type(raw: str) -> str:
    s = raw.strip()
    for pat, qtype in TYPE_RULES:
        if pat.search(s):
            return qtype
    # default conservador
    return TYPE_SCALE


@dataclass
class ParsedQuestion:
    text: str
    question_type: str
    is_required: bool


@dataclass
class ParsedSection:
    name: str
    questions: List[ParsedQuestion]


@dataclass
class ParsedTemplate:
    name: str
    sections: List[ParsedSection]


def parse_docx(docx_path: Path) -> ParsedTemplate:
    if Document is None:
        raise CommandError("python-docx no está disponible en este entorno.")

    doc = Document(str(docx_path))
    template_name: Optional[str] = None
    sections: List[ParsedSection] = []
    current_section: Optional[ParsedSection] = None

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        style = (p.style.name or "").strip() if p.style else ""

        if style == "Heading 1":
            template_name = text
            continue

        if style == "Heading 2":
            current_section = ParsedSection(name=text, questions=[])
            sections.append(current_section)
            continue

        # pregunta (párrafo normal)
        if current_section is None:
            # si no hay Heading 2, creamos sección por defecto
            current_section = ParsedSection(name="General", questions=[])
            sections.append(current_section)

        q = ParsedQuestion(
            text=normalize_question_text(text),
            question_type=detect_question_type(text),
            is_required=detect_is_required(text),
        )
        # evita “preguntas vacías” tras normalización
        if q.text:
            current_section.questions.append(q)

    # --- TABLAS: intentar extraer criterios evaluables ---
    criteria_section = ParsedSection(name="Criterios de evaluación", questions=[])

    for table in doc.tables:
        for row in table.rows:
            cells = [clean_cell_text(cell.text) for cell in row.cells]
            if not any(cells):
                continue

            first = cells[0]
            if not first:
                continue

            # Filtrado conservador
            if looks_like_admin_field(first):
                continue
            if looks_like_header_row(cells):
                continue
            # Evitar líneas tipo "Bloque A – ..." (esto es un título, no un criterio)
            if re.match(r"^bloque\s+[a-e]\b", first.strip().lower()):
                continue

            # Si parece un criterio, lo añadimos (SCALE_1_5)
            w = extract_weight(cells)
            label = first
            # Si hay peso y no está ya en el texto, lo anexamos para inspección en DRY-RUN
            if w and w not in label:
                label = f"{label} ({w})"

            q = ParsedQuestion(
                text=label,
                question_type=TYPE_SCALE,
                is_required=True,   # recomendación: los criterios evaluables deben ser obligatorios
            )
            criteria_section.questions.append(q)

    # Si detectamos criterios por tabla, los añadimos como sección al final
    if criteria_section.questions:
        sections.append(criteria_section)

    if not template_name:
        template_name = docx_path.stem

    # limpia secciones vacías
    sections = [s for s in sections if s.questions]

    return ParsedTemplate(name=template_name, sections=sections)


class Command(BaseCommand):
    help = "Importa (o previsualiza) plantillas de evaluación desde un DOCX."

    def add_arguments(self, parser):
        parser.add_argument("docx", type=str, help="Ruta al fichero .docx")
        parser.add_argument("--apply", action="store_true", help="Escribe en BD (por defecto solo previsualiza).")
        parser.add_argument("--template-code", type=str, default="", help="Código/slug de plantilla (opcional).")

    def handle(self, *args, **options):
        docx = Path(options["docx"])
        if not docx.exists() or docx.suffix.lower() != ".docx":
            raise CommandError(f"Fichero inválido: {docx}")

        parsed = parse_docx(docx)

        # DRY RUN output (siempre lo mostramos)
        self.stdout.write(self.style.SUCCESS(f"Plantilla detectada: {parsed.name}"))
        for i, sec in enumerate(parsed.sections, start=1):
            self.stdout.write(f"  [{i}] Sección: {sec.name}")
            for j, q in enumerate(sec.questions, start=1):
                req = "REQ" if q.is_required else "opt"
                self.stdout.write(f"     - ({q.question_type}, {req}) {q.text}")

        if not options["apply"]:
            self.stdout.write(self.style.WARNING("DRY-RUN: no se ha escrito nada en BD. Usa --apply para importar."))
            return

        base_code = (options["template_code"] or "").strip().upper()
        if not base_code:
            base_code = derive_base_code_from_filename(docx)

        fp = template_fingerprint(parsed)

        existing_same = EvaluationTemplate.objects.filter(base_code=base_code, source_hash=fp).first()
        if existing_same:
            self.stdout.write(
                self.style.WARNING(
                    f"IMPORT SKIP: ya existe {base_code} v{existing_same.version} con el mismo contenido."
                )
            )
            return

        existing_versions = list(
            EvaluationTemplate.objects.filter(base_code=base_code)
            .values_list("version", flat=True)
        )
        max_v = max(existing_versions, default=0)
        new_version = max_v + 1

        with transaction.atomic():
            tpl = EvaluationTemplate.objects.create(
                name=parsed.name,
                base_code=base_code,
                version=new_version,
                source_hash=fp,
                is_active=False,
            )

            for s_idx, sec in enumerate(parsed.sections, start=1):
                sec_obj = TemplateSection.objects.create(
                    template=tpl,
                    title=sec.name,
                    order=s_idx,
                )
                for q_idx, q in enumerate(sec.questions, start=1):
                    TemplateQuestion.objects.create(
                        section=sec_obj,
                        text=q.text,
                        question_type=q.question_type,
                        required=q.is_required,
                        is_required=q.is_required,
                        order=q_idx,
                    )

            if not EvaluationTemplate.objects.filter(base_code=base_code, is_active=True).exists():
                tpl.is_active = True
                tpl.save(update_fields=["is_active"])

        self.stdout.write(self.style.SUCCESS(f"IMPORT OK: {base_code}.v{new_version}"))
